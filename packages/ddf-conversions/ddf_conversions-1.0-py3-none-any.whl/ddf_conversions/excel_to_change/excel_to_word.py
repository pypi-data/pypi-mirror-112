# !/usr/bin/env python
# -*- coding: utf-8 -*-
"""
根据excel指定参数生成word 暂时只支持到AC29列
"""
import cmd
import copy
import os
import sys
import time

import xlrd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


class Order(cmd.Cmd):
    intro = "Excel转Word系统v1.0，输入 help 或者?查看帮助。\n"
    prompt = "command>"

    def do_demo_output(self, arg):
        """样本输出"""
        self.main(is_demo=True)

    def do_complete_output(self, arg):
        """完整输出"""
        self.main(is_demo=False)

    def check_word_file(self, path):
        if path:
            if os.path.splitext(path)[-1] in [".docx", ".doc"]:
                return path
            else:
                raise Exception("同级目录下word格式不正确")
        else:
            if os.path.isfile("demo.docx"):
                return "demo.docx"
            elif os.path.isfile("demo.doc"):
                return "demo.doc"
            else:
                raise Exception("同级目录下不存在word样本文件")

    def check_excel_file(self, path):
        if path:
            if os.path.splitext(path)[-1] in [".xlsx", ".xls"]:
                return path
            else:
                raise Exception("同级目录下excel格式不正确")
        else:
            if os.path.isfile("demo.xlsx"):
                return "demo.xlsx"
            elif os.path.isfile("demo.xls"):
                return "demo.xls"
            else:
                raise Exception("同级目录下不存在excel样本文件")

    def main(self, word_path: str="", excel_path: str="", out_path: str="", is_demo: bool=True):
        try:
            word_path = self.check_word_file(word_path)
            excel_path = self.check_excel_file(excel_path)
            f = ExcelToWord(word_path, excel_path, out_path)
            f.run(is_demo)
        except Exception as e:
            print(e)

    def do_exit(self, _):
        '退出'
        exit(0)


class ExcelToWord(object):
    def __init__(self, word_path, excel_path, out_path=""):
        self.document = Document(word_path)
        self.excel_path = excel_path
        self.out_path = out_path if out_path else ""
        self.format_dict = self.format_to_index()

    def format_to_index(self):
        str_text = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        res = {f"{item}{index + 1}": index for index, item in enumerate(str_text)}
        res.update({"AA27": 26, "AB28": 27, "AC29": 28})
        return res

    def save(self, document, name):
        save_dir = os.getcwd()
        save_path = f"{save_dir}{os.path.sep}dir"
        if os.path.exists(save_path):
            document.save(f"{save_path}{os.path.sep}{name}.docx")
        else:
            os.mkdir(save_path)
            document.save(f"{save_path}{os.path.sep}{name}.docx")

    def save_demo_docx(self, data):
        document = copy.deepcopy(self.document)
        for para in document.paragraphs:
            text = para.text.strip()
            item = text.split(":")
            if text in self.format_dict.keys():
                para.text = str(data[self.format_dict.get(text)])
            # 处理特殊数据
            elif item and item[-1] in self.format_dict.keys():
                para.text = para.text.replace(item[-1], str(data[self.format_dict.get(item[-1])]))
            elif item and item[-1] and item[-1].split("\t")[-1] in self.format_dict.keys():
                para.text = para.text.replace(item[-1].split("\t")[-1],
                                              str(data[self.format_dict.get(item[-1].split("\t")[-1])]))
        # 表格
        tbs = document.tables
        for tb in tbs:
            #行
            for r, row in enumerate(tb.rows):
                #列
                for c, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text in self.format_dict.keys():
                        cell.text = str(data[self.format_dict.get(text)])
                        if text in ("AA27", "AB28", "AC29", "Z26", "X24", "Y25"):
                            tb.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        #保存
        self.save(document, data[0])

    def run(self, is_demo: bool = True):
        sys.stdout.write(f"{ExcelToWord.__name__} start:{int(time.time())}\n")
        excel_data = xlrd.open_workbook(self.excel_path)
        table = excel_data.sheet_by_index(0)
        row = 4 if is_demo else table.nrows
        for i in range(1, row):
            self.save_demo_docx(table.row_values(i))
        sys.stdout.write(f"{ExcelToWord.__name__} stop:{int(time.time())}\n")


if __name__ == "__main__":
    Order().cmdloop()
